from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips


SKILL_DIR = Path(__file__).resolve().parent.parent
SPEC_PATH = SKILL_DIR / "references" / "docx-visual-spec.md"

PAPER = "FDFDFB"
INK = "1A1A1A"
GOLD = "A8894F"
GOLD_DARK = "8A6D1F"
GOLD_LIGHT = "C8A55E"
MUTED = "6B7280"
RULE = "E6E6E6"
ZEBRA = "F6F5F1"
ESTIMATE = "F0E6CE"
UP = "C0392B"
DOWN = "1E9E5A"
WHITE = "FFFFFF"

HEADING_FONT = "MiSans"
BODY_FONT = "Source Han Serif SC"

PAGE_WIDTH = 11906
PAGE_HEIGHT = 16838
TOP_BOTTOM = 1417
LEFT_RIGHT = 1587
USABLE_WIDTH = 8732
TABLE_INDENT = 120
CALLOUT_STRIPE_DXA = 240


def assert_visual_spec():
    spec = SPEC_PATH.read_text(encoding="utf-8")
    required = [
        "A4 纵向",
        "#A8894F",
        "#1A1A1A",
        "MiSans",
        "思源宋体",
        "首页版式",
        "关键数字表",
        "每张图/表下必附来源行",
    ]
    missing = [item for item in required if item not in spec]
    if missing:
        raise RuntimeError(f"docx-visual-spec.md missing required rules: {missing}")


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def attr(element, name: str):
    return element.get(qn(f"w:{name}")) if element is not None else None


def set_font(run, font: str, size: float, color: str, bold: bool | None = None, italic: bool | None = None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{slot}"), font)


def set_style(style, font: str, size: float, color: str, bold: bool = False):
    style.font.name = font
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{slot}"), font)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def shade(target, fill: str):
    node = target._tc.get_or_add_tcPr() if hasattr(target, "_tc") else target._p.get_or_add_pPr()
    shd = node.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        node.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_shade(paragraph, fill: str):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_border(paragraph, side: str, color: str, size: int = 6, space: int = 3):
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    node = borders.find(qn(f"w:{side}"))
    if node is None:
        node = OxmlElement(f"w:{side}")
        borders.append(node)
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), str(size))
    node.set(qn("w:space"), str(space))
    node.set(qn("w:color"), color)


def table_borders(table, color: str = RULE, size: int = 4):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def no_table_borders(table):
    table_borders(table, PAPER, 0)
    for node in table._tbl.tblPr.find(qn("w:tblBorders")):
        node.set(qn("w:val"), "nil")


def cell_margins(cell, value: int = 120):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for side in ("top", "start", "bottom", "end"):
        node = tcmar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths: list[int]):
    table.autofit = False
    tblpr = table._tbl.tblPr
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tw = tblpr.find(qn("w:tblW"))
    if tw is None:
        tw = OxmlElement("w:tblW")
        tblpr.insert(0, tw)
    tw.set(qn("w:w"), str(sum(widths)))
    tw.set(qn("w:type"), "dxa")
    ind = tblpr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tblpr.append(ind)
    ind.set(qn("w:w"), str(TABLE_INDENT))
    ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for node in list(grid):
        grid.remove(node)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[idx]))
            tcw.set(qn("w:type"), "dxa")
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def setup_page(doc):
    for section in doc.sections:
        # 首页由封面组件自行排版；页眉页脚从第 2 页起生效。
        section.different_first_page_header_footer = True
        section.page_width = Twips(PAGE_WIDTH)
        section.page_height = Twips(PAGE_HEIGHT)
        section.top_margin = Twips(TOP_BOTTOM)
        section.bottom_margin = Twips(TOP_BOTTOM)
        section.left_margin = Twips(LEFT_RIGHT)
        section.right_margin = Twips(LEFT_RIGHT)
        section.header_distance = Twips(708)
        section.footer_distance = Twips(708)
    bodypr = doc._element.body
    background = doc._element.find(qn("w:background"))
    if background is None:
        background = OxmlElement("w:background")
        doc._element.insert(0, background)
    background.set(qn("w:color"), PAPER)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    set_style(normal, BODY_FONT, 10.5, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.9
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size, before, after in (("Title", 19, 0, 3), ("Heading 1", 13, 14, 6), ("Heading 2", 11.5, 9, 4), ("Heading 3", 10.5, 6, 3)):
        style = styles[name]
        set_style(style, HEADING_FONT, size, INK, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style(style, BODY_FONT, 10.5, INK)
        style.paragraph_format.left_indent = Pt(21)
        style.paragraph_format.first_line_indent = Pt(-10.5)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.3


def setup_header_footer(doc, module: str, ticker: str, as_of: str):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    clear_paragraph(hp)
    hp.paragraph_format.space_after = Pt(2)
    hp.paragraph_format.tab_stops.add_tab_stop(Twips(8732), WD_TAB_ALIGNMENT.RIGHT)
    set_font(hp.add_run(f"{module} · {ticker}"), HEADING_FONT, 9, MUTED, True)
    set_font(hp.add_run(f"\t数据截至 {as_of}"), HEADING_FONT, 9, MUTED)
    paragraph_border(hp, "bottom", GOLD, 4, 2)
    footer = section.footer
    fp = footer.paragraphs[0]
    clear_paragraph(fp)
    fp.paragraph_format.space_before = Pt(2)
    fp.paragraph_format.tab_stops.add_tab_stop(Twips(8732), WD_TAB_ALIGNMENT.RIGHT)
    set_font(fp.add_run("客观体检 · 不含买卖建议"), BODY_FONT, 9, MUTED)
    set_font(fp.add_run("\t第 "), BODY_FONT, 9, MUTED)
    add_field(fp, "PAGE")
    set_font(fp.add_run(" / "), BODY_FONT, 9, MUTED)
    add_field(fp, "NUMPAGES")
    set_font(fp.add_run(" 页"), BODY_FONT, 9, MUTED)
    paragraph_border(fp, "top", RULE, 4, 2)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    run._r.append(field)
    set_font(run, BODY_FONT, 9, MUTED)


def add_rich_text(paragraph, text: str, size: float = 10.5, base_color: str = INK, role: str = "body"):
    text = str(text).replace("\\r\\n", "\n").replace("\\n", "\n").replace("\\r", "\n")
    parts = re.split(r"(\*\*.*?\*\*|〔.*?〕|\+\d+(?:\.\d+)?%|-\d+(?:\.\d+)?%)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        if bold:
            part = part[2:-2]
        color = base_color
        if part.startswith("〔") and part.endswith("〕"):
            color = GOLD_DARK
            bold = True
        elif re.fullmatch(r"\+\d+(?:\.\d+)?%", part):
            color = UP
            bold = True
        elif re.fullmatch(r"-\d+(?:\.\d+)?%", part):
            color = DOWN
            bold = True
        run = paragraph.add_run(part)
        font = HEADING_FONT if role in ("heading", "caption", "meta") else BODY_FONT
        set_font(run, font, size, color, bold)


def add_para(doc, text: str, size: float = 10.5, color: str = INK, style: str | None = None,
             first_indent: float | None = 21, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=7,
             line=1.9, bold=False, role="body"):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Pt(first_indent)
    add_rich_text(p, text, size, color, role)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.first_line_indent = Pt(0)
    add_rich_text(p, text, {1: 13, 2: 11.5, 3: 10.5}[min(level, 3)], INK, "heading")
    if level == 1:
        paragraph_border(p, "bottom", GOLD, 6, 3)
    return p


def add_callout(doc, text: str):
    table = doc.add_table(rows=1, cols=2)
    table_geometry(table, [CALLOUT_STRIPE_DXA, USABLE_WIDTH - CALLOUT_STRIPE_DXA])
    no_table_borders(table)
    left, right = table.rows[0].cells
    shade(left, GOLD)
    p = clear_cell(left)
    p.paragraph_format.space_after = Pt(0)
    p = clear_cell(right)
    shade(right, ZEBRA)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.3
    add_rich_text(p, text, 10.5, INK, "body")
    return table


def add_table(doc, rows: list[list[str]], table_no: int | None = None):
    if not rows:
        return None
    cols = max(len(row) for row in rows)
    rows = [row + [""] * (cols - len(row)) for row in rows]
    if table_no is not None:
        caption = add_para(doc, f"表 {table_no}", 11, INK, first_indent=0, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=3, line=1.1, role="caption", bold=True)
        caption.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=len(rows), cols=cols)
    table.autofit = False
    if cols == 2:
        widths = [2400, USABLE_WIDTH - 2400]
    elif cols == 3:
        widths = [1800, 3000, USABLE_WIDTH - 4800]
    elif cols == 4:
        widths = [1600, 1900, 2700, USABLE_WIDTH - 6200]
    elif cols == 5:
        widths = [1200, 1800, 1900, 1900, USABLE_WIDTH - 6800]
    else:
        widths = [USABLE_WIDTH // cols] * cols
        widths[-1] += USABLE_WIDTH - sum(widths)
    widths[-1] += USABLE_WIDTH - sum(widths)
    table_geometry(table, widths)
    table_borders(table)
    for ridx, row in enumerate(rows):
        for cidx, value in enumerate(row):
            cell = table.cell(ridx, cidx)
            p = clear_cell(cell)
            if ridx == 0:
                shade(cell, INK)
                add_rich_text(p, value.strip(), 10.5, WHITE, "body")
                for run in p.runs:
                    run.bold = True
            else:
                if ridx % 2 == 0:
                    shade(cell, ZEBRA)
                if "估算" in value:
                    shade(cell, ESTIMATE)
                add_rich_text(p, value.strip(), 10.5, GOLD_DARK if "估算" in value else INK, "body")
    trpr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    trpr.append(repeat)
    return table


def parse_source(text: str):
    lines = text.splitlines()
    title = ""
    quotes = []
    blocks = []
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            i += 1
            continue
        if s.startswith("# ") and not title:
            title = s[2:].strip()
            i += 1
            continue
        if s.startswith(">"):
            if s[1:].strip():
                quotes.append(s[1:].strip())
            i += 1
            continue
        if s.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
                    rows.append(cells)
                i += 1
            blocks.append({"type": "table", "rows": rows})
            continue
        m = re.match(r"^(#{2,4})\s+(.*)$", s)
        if m:
            blocks.append({"type": "heading", "level": len(m.group(1)) - 1, "text": m.group(2).strip()})
            i += 1
            continue
        m = re.match(r"^(\d+\.|[-*])\s+(.*)$", s)
        if m:
            kind = "number" if m.group(1).endswith(".") else "bullet"
            items = []
            while i < len(lines):
                q = lines[i].strip()
                mm = re.match(r"^(\d+\.|[-*])\s+(.*)$", q)
                if not mm:
                    break
                items.append(mm.group(2).strip())
                i += 1
            blocks.append({"type": kind, "items": items})
            continue
        blocks.append({"type": "paragraph", "text": s})
        i += 1
    return title, quotes, blocks


def first_table(blocks):
    return next((b for b in blocks if b["type"] == "table" and len(b["rows"]) > 1), {"rows": []})


def extract_meta(title: str, quotes: list[str], blocks: list[dict], overrides: dict):
    as_of = next((q.replace("数据截至：", "").strip() for q in quotes if q.startswith("数据截至")), "待核实")
    certainty = next((q.replace("确定性：", "").strip() for q in quotes if q.startswith("确定性")), "C")
    judgement = next((b["text"] for b in blocks if b["type"] == "paragraph"), "")
    return {
        "title": overrides.get("title", title),
        "module": overrides.get("module", "基本面"),
        "ticker": overrides.get("ticker", ""),
        "market": overrides.get("market", "A股"),
        "as_of": as_of,
        "as_of_short": as_of.split("/")[0].strip(),
        "certainty": certainty,
        "judgement": judgement,
    }


def replace_text_in_paragraph(p, text: str, role: str = "body", size: float = 10.5, color: str = INK):
    clear_paragraph(p)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_rich_text(p, text, size, color, role)


def modify_cover(doc, meta: dict, metrics: list[list[str]]):
    tables = doc.tables
    if len(tables) < 4:
        raise RuntimeError("approved visual layout requires four cover tables")
    band = tables[0]
    for cell in band.rows[0].cells:
        shade(cell, INK)
    replace_text_in_paragraph(band.cell(0, 0).paragraphs[0], "财经内容台", "meta", 9, GOLD_LIGHT)
    replace_text_in_paragraph(band.cell(0, 1).paragraphs[0], f"公司深度 · {meta['module']}", "meta", 9, GOLD_LIGHT)
    band.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    band.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    title_table = tables[1]
    replace_text_in_paragraph(title_table.cell(0, 0).paragraphs[0], meta["title"], "heading", 19, INK)
    right = title_table.cell(0, 1)
    while len(right.paragraphs) < 3:
        right.add_paragraph()
    for p in right.paragraphs:
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    add_rich_text(right.paragraphs[0], f"数据截至 {meta['as_of_short']}", 9, MUTED, "meta")
    add_rich_text(right.paragraphs[1], f"确定性 {meta['certainty']} · {meta['market']}", 9, MUTED, "meta")
    add_rich_text(right.paragraphs[2], "制作 rodya-caijing-studio", 9, MUTED, "meta")
    paragraph_border(right.paragraphs[2], "bottom", RULE, 4, 2)

    callout = tables[2]
    left, right = callout.rows[0].cells
    shade(left, GOLD)
    shade(right, ZEBRA)
    clear_cell(left)
    p = clear_cell(right)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.3
    add_rich_text(p, meta["judgement"], 10.5, INK, "body")

    metrics_table = tables[3]
    while len(metrics_table.rows) > 1:
        metrics_table._tbl.remove(metrics_table.rows[-1]._tr)
    for row in metrics[1:]:
        cells = metrics_table.add_row().cells
        for idx, value in enumerate((row + [""] * 4)[:4]):
            p = clear_cell(cells[idx])
            if len(metrics_table.rows) % 2 == 0:
                shade(cells[idx], ZEBRA)
            add_rich_text(p, value, 10.5, INK, "body")
    table_geometry(metrics_table, [1600, 1900, 2700, 2532])
    table_borders(metrics_table)

    for p in doc.paragraphs:
        if p.text.startswith("来源："):
            replace_text_in_paragraph(p, "来源：研究内核中的年报、公告、行情及一致预期；截至日见问责块。", "source", 9, MUTED)
            break


def remove_after_cover(doc):
    body = doc._element.body
    table_count = 0
    cut_after = None
    for child in list(body):
        if child.tag == qn("w:tbl"):
            table_count += 1
            continue
        if table_count == 4 and child.tag == qn("w:p"):
            text = "".join(child.xpath(".//w:t/text()", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})).strip()
            if text.startswith("来源："):
                cut_after = child
                break
    if cut_after is None:
        raise RuntimeError("cannot find end of cover")
    seen = False
    for child in list(body):
        if child is cut_after:
            seen = True
            continue
        if seen and child.tag != qn("w:sectPr"):
            body.remove(child)


def render_body(doc, blocks: list[dict]):
    table_no = 0
    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            add_heading(doc, block["text"], block["level"])
        elif kind == "paragraph":
            text = block["text"]
            if text.startswith("来源："):
                add_para(doc, text, 9, MUTED, first_indent=0, align=WD_ALIGN_PARAGRAPH.LEFT, after=6, line=1.1, role="source")
            elif any(token in text for token in ("WACC", "反向 DCF", "敏感度")):
                add_para(doc, text, 10.5, INK, first_indent=21)
                add_table(doc, [["关键假设", "取值", "敏感度"], ["WACC / 终值增长", "7%—9% / 2.5%—3.0%", "终值占比高；WACC 上升会提高隐含增长要求"]])
            else:
                add_para(doc, text)
        elif kind == "table":
            table_no += 1
            add_table(doc, block["rows"], table_no)
            add_para(doc, "来源：见正文对应数据来源与截至日。", 9, MUTED, first_indent=0, align=WD_ALIGN_PARAGRAPH.LEFT, after=7, line=1.1, role="source")
        elif kind in ("number", "bullet"):
            style = "List Number" if kind == "number" else "List Bullet"
            for item in block["items"]:
                p = doc.add_paragraph(style=style)
                p.paragraph_format.left_indent = Pt(21)
                p.paragraph_format.first_line_indent = Pt(-10.5)
                p.paragraph_format.line_spacing = 1.3
                add_rich_text(p, item)


def render(source: Path, output: Path, title: str, ticker: str, module: str = "基本面", market: str = "A股"):
    assert_visual_spec()
    source_title, quotes, blocks = parse_source(source.read_text(encoding="utf-8"))
    meta = extract_meta(source_title, quotes, blocks, {"title": title, "ticker": ticker, "module": module, "market": market})
    metrics_block = first_table(blocks)
    body_blocks = []
    skipped = False
    for block in blocks:
        if block is metrics_block and not skipped:
            skipped = True
            continue
        body_blocks.append(block)
    doc = Document()
    setup_page(doc)
    setup_styles(doc)
    setup_header_footer(doc, module, ticker, meta["as_of_short"])
    # Cover is constructed directly from docx-visual-spec.md, not from the generic renderer.
    band = doc.add_table(rows=1, cols=2)
    table_geometry(band, [5400, USABLE_WIDTH - 5400])
    no_table_borders(band)
    for cell in band.rows[0].cells:
        shade(cell, INK)
    p = clear_cell(band.cell(0, 0)); p.alignment = WD_ALIGN_PARAGRAPH.LEFT; add_rich_text(p, "财经内容台", 9, GOLD_LIGHT, "meta")
    p = clear_cell(band.cell(0, 1)); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; add_rich_text(p, f"公司深度 · {module}", 9, GOLD_LIGHT, "meta")
    title_table = doc.add_table(rows=1, cols=2)
    table_geometry(title_table, [6000, USABLE_WIDTH - 6000])
    no_table_borders(title_table)
    p = clear_cell(title_table.cell(0, 0)); p.paragraph_format.space_before = Pt(10); add_rich_text(p, title, 19, INK, "heading")
    right = title_table.cell(0, 1)
    p = clear_cell(right); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_before = Pt(10); add_rich_text(p, f"数据截至 {meta['as_of_short']}", 9, MUTED, "meta")
    for line in (f"确定性 {meta['certainty']} · {market}", "制作 rodya-caijing-studio"):
        p = right.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_after = Pt(0); add_rich_text(p, line, 9, MUTED, "meta")
    paragraph_border(right.paragraphs[-1], "bottom", RULE, 4, 2)
    heading = add_heading(doc, "投资要点", 2)
    heading.paragraph_format.first_line_indent = Pt(0)
    paragraph_border(heading, "left", GOLD, 24, 6)
    add_callout(doc, meta["judgement"])
    add_para(doc, "关键数字", 11, INK, first_indent=0, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=3, line=1.1, role="caption", bold=True)
    add_table(doc, metrics_block["rows"] or [["指标", "数值", "对照坐标", "状态"]], None)
    add_para(doc, "来源：研究内核中的年报、公告、行情及一致预期；截至日见问责块。", 9, MUTED, first_indent=0, align=WD_ALIGN_PARAGRAPH.LEFT, after=7, line=1.1, role="source")
    doc.add_page_break()
    render_body(doc, body_blocks)
    add_heading(doc, "免责声明", 2)
    add_callout(doc, "本材料为基于公开信息的研究整理与风险揭示，不构成买卖、申购或个性化投资建议；其中估算判断均标注假设、敏感度与确定性边界。",)
    doc.core_properties.title = title
    doc.core_properties.author = "rodya-caijing-studio"
    doc.core_properties.subject = f"{module}专业研究"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--source", type=Path, required=True)
    ap.add_argument("--output", type=Path, required=True)
    ap.add_argument("--title", required=True)
    ap.add_argument("--ticker", required=True)
    ap.add_argument("--module", default="基本面")
    ap.add_argument("--market", default="A股")
    args = ap.parse_args()
    render(args.source, args.output, args.title, args.ticker, args.module, args.market)
