from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
TOKENS_PATH = SKILL_DIR / "references" / "docx-style-tokens.json"
TEMPLATE_PATH = SKILL_DIR / "assets" / "docx-template.docx"


def load_tokens(path: Path = TOKENS_PATH) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def wval(element, attribute="val"):
    if element is None:
        return None
    return element.get(qn(f"w:{attribute}"))


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def normalize_docx_text(text: str) -> str:
    """Turn escaped line-break markers from generated content into real breaks."""
    return str(text).replace("\\r\\n", "\n").replace("\\n", "\n").replace("\\r", "\n")


def set_run_font(run, font: str, fallback: str, size: float, color: str,
                 bold: bool | None = None, italic: bool | None = None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font if attr == "eastAsia" else fallback if attr in ("ascii", "hAnsi") else font)


def set_style_font(style, font: str, fallback: str, size: float, color: str, bold: bool = False):
    style.font.name = font
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), fallback)
    rfonts.set(qn("w:hAnsi"), fallback)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:cs"), font)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, value: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "E6E6E6", size: int = 4):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_table_geometry(table, widths: list[int], tokens: dict):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(tokens["page"]["table_indent_dxa"]))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, tokens["table"]["cell_margin_dxa"])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_border(paragraph, side: str, color: str, size: int = 6, space: int = 4):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    node = borders.find(qn(f"w:{side}"))
    if node is None:
        node = OxmlElement(f"w:{side}")
        borders.append(node)
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), str(size))
    node.set(qn("w:space"), str(space))
    node.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_field(paragraph, instruction: str, tokens: dict):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    run._r.append(fld)
    set_run_font(run, tokens["fonts"]["source"], tokens["fonts"]["body_fallback"], tokens["sizes_pt"]["source"], tokens["colors"]["muted"])


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def apply_styles(doc, tokens: dict):
    fonts, colors, sizes = tokens["fonts"], tokens["colors"], tokens["sizes_pt"]
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, fonts["body"], fonts["body_fallback"], sizes["body"], colors["ink"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = tokens["body"]["line_spacing"]
    normal.paragraph_format.first_line_indent = Pt(tokens["body"]["first_line_indent_pt"])
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size, before, after in (("Title", sizes["title"], 0, 3), ("Heading 1", sizes["h1"], 14, 6), ("Heading 2", sizes["h2"], 9, 4), ("Heading 3", sizes["h3"], 6, 3)):
        style = styles[name]
        set_style_font(style, fonts["heading"], fonts["heading_fallback"], size, colors["ink"], True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style_font(style, fonts["body"], fonts["body_fallback"], sizes["body"], colors["ink"])
        style.paragraph_format.left_indent = Pt(21)
        style.paragraph_format.first_line_indent = Pt(-10.5)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.3


def set_page_setup(section, tokens: dict):
    page = tokens["page"]
    section.page_width = Twips(page["width_dxa"])
    section.page_height = Twips(page["height_dxa"])
    section.top_margin = Twips(page["top_margin_dxa"])
    section.bottom_margin = Twips(page["bottom_margin_dxa"])
    section.left_margin = Twips(page["left_margin_dxa"])
    section.right_margin = Twips(page["right_margin_dxa"])
    section.header_distance = Twips(567)
    section.footer_distance = Twips(567)
    section.different_first_page_header_footer = True


def set_header_footer(section, metadata: dict, tokens: dict):
    colors, fonts, sizes = tokens["colors"], tokens["fonts"], tokens["sizes_pt"]
    header = section.header
    hp = header.paragraphs[0]
    clear_paragraph(hp)
    hp.paragraph_format.space_after = Pt(2)
    hp.paragraph_format.tab_stops.add_tab_stop(Pt(426), WD_TAB_ALIGNMENT.RIGHT)
    r = hp.add_run(f"{metadata['module']} · {metadata['ticker']}")
    set_run_font(r, fonts["heading"], fonts["heading_fallback"], sizes["meta"], colors["muted"], True)
    r = hp.add_run(f"\t数据截至 {metadata['as_of_short']}")
    set_run_font(r, fonts["heading"], fonts["heading_fallback"], sizes["meta"], colors["muted"])
    set_paragraph_border(hp, "bottom", colors["gold"], 4, 2)
    footer = section.footer
    fp = footer.paragraphs[0]
    clear_paragraph(fp)
    fp.paragraph_format.space_before = Pt(2)
    fp.paragraph_format.tab_stops.add_tab_stop(Pt(426), WD_TAB_ALIGNMENT.RIGHT)
    r = fp.add_run("客观体检 · 不含买卖建议")
    set_run_font(r, fonts["body"], fonts["body_fallback"], sizes["source"], colors["muted"])
    r = fp.add_run("\t第 ")
    set_run_font(r, fonts["body"], fonts["body_fallback"], sizes["source"], colors["muted"])
    add_field(fp, "PAGE", tokens)
    r = fp.add_run(" / ")
    set_run_font(r, fonts["body"], fonts["body_fallback"], sizes["source"], colors["muted"])
    add_field(fp, "NUMPAGES", tokens)
    r = fp.add_run(" 页")
    set_run_font(r, fonts["body"], fonts["body_fallback"], sizes["source"], colors["muted"])
    set_paragraph_border(fp, "top", colors["rule"], 4, 2)


def add_text(paragraph, text: str, tokens: dict, role: str = "body", bold: bool = False,
             color: str | None = None):
    fonts, colors, sizes = tokens["fonts"], tokens["colors"], tokens["sizes_pt"]
    font = fonts["heading"] if role in ("heading", "meta") else fonts["body"]
    fallback = fonts["heading_fallback"] if role in ("heading", "meta") else fonts["body_fallback"]
    size = sizes[{"heading": "h2", "meta": "meta", "body": "body", "table": "table", "source": "source", "title": "title"}[role]]
    run = paragraph.add_run(normalize_docx_text(text))
    set_run_font(run, font, fallback, size, color or colors["ink"], bold=bold)
    return run


def add_body_paragraph(doc, text: str, tokens: dict, style: str = "Normal"):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if style == "Normal":
        p.paragraph_format.first_line_indent = Pt(tokens["body"]["first_line_indent_pt"])
        p.paragraph_format.line_spacing = tokens["body"]["line_spacing"]
    add_text(p, text, tokens, "body")
    return p


def add_heading(doc, text: str, level: int, tokens: dict):
    style = f"Heading {min(level, 3)}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True
    add_text(p, text, tokens, "heading", True)
    if level == 1:
        set_paragraph_border(p, "bottom", tokens["colors"]["gold"], 6, 3)
    return p


def add_callout(doc, text: str, tokens: dict):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Pt(10)
    p.paragraph_format.right_indent = Pt(5)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.3
    set_paragraph_shading(p, tokens["colors"]["estimate_fill"])
    set_paragraph_border(p, "left", tokens["colors"]["gold"], 18, 6)
    add_text(p, text, tokens, "body")
    return p


def add_table(doc, rows: list[list[str]], tokens: dict, header: bool = True):
    if not rows:
        return None
    cols = max(len(r) for r in rows)
    rows = [r + [""] * (cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=cols)
    table.autofit = False
    width = tokens["page"]["usable_width_dxa"]
    if cols == 2:
        widths = [int(width * 0.26), width - int(width * 0.26)]
    elif cols == 3:
        widths = [int(width * 0.21), int(width * 0.32), width - int(width * 0.53)]
    elif cols == 4:
        widths = [int(width * 0.18), int(width * 0.22), int(width * 0.30), width - int(width * 0.70)]
    elif cols == 5:
        widths = [int(width * 0.14), int(width * 0.20), int(width * 0.23), int(width * 0.21), width - int(width * 0.78)]
    else:
        widths = [width // cols] * cols
        widths[-1] += width - sum(widths)
    widths[-1] += width - sum(widths)
    set_table_geometry(table, widths, tokens)
    set_table_borders(table, tokens["colors"]["rule"], 4)
    for ridx, row in enumerate(rows):
        for cidx, value in enumerate(row):
            cell = table.cell(ridx, cidx)
            p = clear_cell(cell)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if cidx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            if ridx == 0 and header:
                set_cell_shading(cell, tokens["colors"]["ink"])
                add_text(p, value.strip(), tokens, "table", True, tokens["colors"]["white"])
            else:
                if ridx % 2 == 0:
                    set_cell_shading(cell, tokens["colors"]["zebra"])
                add_text(p, value.strip(), tokens, "table")
    if header:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        header_node = OxmlElement("w:tblHeader")
        header_node.set(qn("w:val"), "true")
        tr_pr.append(header_node)
    return table


def parse_markdown(text: str):
    lines = text.splitlines()
    blocks = []
    title = ""
    quotes = []
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# ") and not title:
            title = stripped[2:].strip()
            i += 1
            continue
        if stripped.startswith(">"):
            quote = stripped[1:].strip()
            if quote:
                quotes.append(quote)
            i += 1
            continue
        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
                    rows.append(cells)
                i += 1
            blocks.append({"type": "table", "rows": rows})
            continue
        match = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if match:
            blocks.append({"type": "heading", "level": len(match.group(1)) - 1, "text": match.group(2).strip()})
            i += 1
            continue
        match = re.match(r"^(\d+\.|[-*])\s+(.*)$", stripped)
        if match:
            kind = "number" if match.group(1).endswith(".") else "bullet"
            items = []
            while i < len(lines):
                current = lines[i].strip()
                m = re.match(r"^(\d+\.|[-*])\s+(.*)$", current)
                if not m:
                    break
                items.append(m.group(2).strip())
                i += 1
            blocks.append({"type": kind, "items": items})
            continue
        blocks.append({"type": "paragraph", "text": stripped})
        i += 1
    return {"title": title, "quotes": quotes, "blocks": blocks}


def extract_metadata(parsed: dict, overrides: dict):
    title = overrides.get("title") or parsed["title"]
    data_as_of = next((x for x in parsed["quotes"] if x.startswith("数据截至")), "数据截至：待核实")
    certainty = next((x for x in parsed["quotes"] if x.startswith("确定性")), "确定性：C")
    as_of = data_as_of.replace("数据截至：", "").strip()
    certainty_value = certainty.replace("确定性：", "").strip()
    first_judgement = ""
    for block in parsed["blocks"]:
        if block["type"] == "paragraph":
            first_judgement = block["text"]
            break
    metrics = next((b for b in parsed["blocks"] if b["type"] == "table" and len(b["rows"]) >= 2), {"rows": []})
    body_blocks = []
    skipped_metrics = False
    for block in parsed["blocks"]:
        if block is metrics and not skipped_metrics:
            skipped_metrics = True
            continue
        body_blocks.append(block)
    return {
        "title": title,
        "module": overrides.get("module", "基本面"),
        "ticker": overrides.get("ticker", ""),
        "market": overrides.get("market", "A股"),
        "skill": overrides.get("skill", "rodya-caijing-studio"),
        "as_of": as_of,
        "as_of_short": as_of.split("/")[0].strip(),
        "certainty": certainty_value,
        "judgement": first_judgement,
        "metrics": metrics["rows"],
        "metric_source": overrides.get("metric_source", "来源：研究内核中的年报、公告、行情及一致预期；截至日见问责块。"),
        "body_blocks": body_blocks,
    }


def replace_placeholders(doc, replacements: dict):
    def replace_paragraph(p):
        if not p.runs:
            return
        text = normalize_docx_text(p.text)
        if not any(key in text for key in replacements):
            return
        new_text = text
        for key, value in replacements.items():
            new_text = new_text.replace(key, normalize_docx_text(str(value)))
        p.runs[0].text = new_text
        for run in p.runs[1:]:
            run.text = ""
    for p in doc.paragraphs:
        replace_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_paragraph(p)


def prepare_cover_metrics(doc, metadata: dict, tokens: dict):
    if len(doc.tables) < 3:
        return
    table = doc.tables[2]
    rows = metadata["metrics"] or [["指标", "数值", "对比", "状态"], ["待补充", "—", "—", "待核实"]]
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row[:4]):
            p = clear_cell(cells[idx])
            if len(row) < 4:
                value = value or "—"
            add_text(p, value, tokens, "table")
            if len(table.rows) % 2 == 0:
                set_cell_shading(cells[idx], tokens["colors"]["zebra"])
    widths = [int(tokens["page"]["usable_width_dxa"] * x) for x in (0.18, 0.22, 0.30, 0.30)]
    widths[-1] += tokens["page"]["usable_width_dxa"] - sum(widths)
    set_table_geometry(table, widths, tokens)


def render_body(doc, blocks: Iterable[dict], tokens: dict):
    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            add_heading(doc, block["text"], block["level"], tokens)
        elif kind == "paragraph":
            add_body_paragraph(doc, block["text"], tokens)
        elif kind == "table":
            add_table(doc, block["rows"], tokens)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
        elif kind in ("number", "bullet"):
            style = "List Number" if kind == "number" else "List Bullet"
            for item in block["items"]:
                p = doc.add_paragraph(style=style)
                p.paragraph_format.first_line_indent = Pt(-10.5)
                p.paragraph_format.left_indent = Pt(21)
                add_text(p, item, tokens, "body")


def render_markdown(source: Path, output: Path, template: Path = TEMPLATE_PATH, overrides: dict | None = None):
    tokens = load_tokens()
    parsed = parse_markdown(source.read_text(encoding="utf-8"))
    metadata = extract_metadata(parsed, overrides or {})
    doc = Document(template)
    for section in doc.sections:
        set_page_setup(section, tokens)
        set_header_footer(section, metadata, tokens)
    apply_styles(doc, tokens)
    replacements = {
        "{{ORG}}": "财经内容台",
        "{{BAND_LABEL}}": f"公司深度 · {metadata['module']}",
        "{{TITLE}}": metadata["title"],
        "{{AS_OF}}": metadata["as_of"],
        "{{CERTAINTY}}": metadata["certainty"],
        "{{MARKET}}": metadata["market"],
        "{{SKILL}}": metadata["skill"],
        "{{JUDGEMENT}}": metadata["judgement"],
        "{{METRIC_SOURCE}}": metadata["metric_source"],
    }
    replace_placeholders(doc, replacements)
    prepare_cover_metrics(doc, metadata, tokens)
    for p in list(doc.paragraphs):
        if "{{BODY_START}}" in p.text:
            p._element.getparent().remove(p._element)
    render_body(doc, metadata["body_blocks"], tokens)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = metadata["title"]
    doc.core_properties.author = "财经内容台"
    doc.core_properties.subject = f"{metadata['module']}专业研究"
    doc.save(output)
    return output


def create_template(path: Path = TEMPLATE_PATH):
    tokens = load_tokens()
    fonts, colors, sizes = tokens["fonts"], tokens["colors"], tokens["sizes_pt"]
    doc = Document()
    apply_styles(doc, tokens)
    section = doc.sections[0]
    set_page_setup(section, tokens)
    set_header_footer(section, {"module": "基本面", "ticker": "{{TICKER}}", "as_of_short": "{{AS_OF}}"}, tokens)
    first = section.first_page_header
    clear_paragraph(first.paragraphs[0])
    first_footer = section.first_page_footer
    clear_paragraph(first_footer.paragraphs[0])

    band = doc.add_table(rows=1, cols=2)
    set_table_geometry(band, [5400, tokens["page"]["usable_width_dxa"] - 5400], tokens)
    remove_table_borders(band)
    for cell in band.rows[0].cells:
        set_cell_shading(cell, colors["ink"])
        p = clear_cell(cell)
        p.paragraph_format.space_after = Pt(0)
    add_text(band.cell(0, 0).paragraphs[0], "{{ORG}}", tokens, "meta", True, colors["gold_on_dark"])
    band.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(band.cell(0, 1).paragraphs[0], "{{BAND_LABEL}}", tokens, "meta", color=colors["gold_on_dark"])
    band.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    title_table = doc.add_table(rows=1, cols=2)
    set_table_geometry(title_table, [6000, tokens["page"]["usable_width_dxa"] - 6000], tokens)
    remove_table_borders(title_table)
    left, right = title_table.rows[0].cells
    p = clear_cell(left)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    add_text(p, "{{TITLE}}", tokens, "title", True)
    rp = clear_cell(right)
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(10)
    rp.paragraph_format.space_after = Pt(2)
    add_text(rp, "数据截至 {{AS_OF}}\n确定性 {{CERTAINTY}} · {{MARKET}}\n制作 {{SKILL}}", tokens, "meta")
    set_paragraph_border(rp, "bottom", colors["rule"], 4, 2)

    heading = doc.add_paragraph(style="Heading 2")
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.first_line_indent = Pt(0)
    set_paragraph_border(heading, "left", colors["gold"], 24, 6)
    add_text(heading, "投资要点", tokens, "heading", True)
    callout = doc.add_paragraph(style="Normal")
    callout.paragraph_format.left_indent = Pt(10)
    callout.paragraph_format.right_indent = Pt(5)
    callout.paragraph_format.first_line_indent = Pt(0)
    callout.paragraph_format.line_spacing = 1.3
    set_paragraph_shading(callout, colors["estimate_fill"])
    set_paragraph_border(callout, "left", colors["gold"], 18, 6)
    add_text(callout, "{{JUDGEMENT}}", tokens, "body")

    label = doc.add_paragraph(style="Heading 2")
    label.paragraph_format.first_line_indent = Pt(0)
    add_text(label, "关键数字", tokens, "heading", True)
    metrics = doc.add_table(rows=2, cols=4)
    set_table_geometry(metrics, [1572, 1921, 2620, 2619], tokens)
    set_table_borders(metrics, colors["rule"], 4)
    headers = ["指标", "数值", "对比 / 口径", "状态"]
    for i, value in enumerate(headers):
        p = clear_cell(metrics.cell(0, i))
        set_cell_shading(metrics.cell(0, i), colors["ink"])
        add_text(p, value, tokens, "table", True)
        p.runs[0].font.color.rgb = rgb(colors["white"])
    for i in range(4):
        p = clear_cell(metrics.cell(1, i))
        add_text(p, "—", tokens, "table")
    source = doc.add_paragraph()
    source.paragraph_format.first_line_indent = Pt(0)
    source.paragraph_format.space_before = Pt(2)
    source.paragraph_format.space_after = Pt(8)
    add_text(source, "{{METRIC_SOURCE}}", tokens, "source")
    page_break = doc.add_paragraph()
    page_break.add_run().add_break()
    marker = doc.add_paragraph()
    add_text(marker, "{{BODY_START}}", tokens, "body")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--create-template", type=Path)
    parser.add_argument("--source", type=Path)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--template", type=Path, default=TEMPLATE_PATH)
    parser.add_argument("--title")
    parser.add_argument("--module", default="基本面")
    parser.add_argument("--ticker", default="")
    parser.add_argument("--market", default="A股")
    args = parser.parse_args()
    if args.create_template:
        create_template(args.create_template)
        return
    if not args.source or not args.output:
        parser.error("rendering requires --source and --output")
    render_markdown(args.source, args.output, args.template, {
        "title": args.title,
        "module": args.module,
        "ticker": args.ticker,
        "market": args.market,
    })


if __name__ == "__main__":
    main()
