from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Twips


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
HEADING_STYLE_RE = re.compile(r"^Heading ([1-3])$")
CHAPTER_TITLE_RE = re.compile(r"^[零一二三四五六七八九十百千万]+、")


def read_tokens(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def require(condition: bool, message: str):
    if not condition:
        raise AssertionError(message)


def audit_heading_hierarchy(doc):
    headings = []
    for paragraph in doc.paragraphs:
        match = HEADING_STYLE_RE.fullmatch(paragraph.style.name or "")
        if not match:
            continue
        level = int(match.group(1))
        text = paragraph.text.strip()
        require(text, f"Heading {level} is empty")
        headings.append((level, text))

    require(headings, "missing Heading 1-3 paragraphs")
    require(any(level == 1 for level, _ in headings), "missing Heading 1 chapter")
    require(any(level == 2 for level, _ in headings), "missing Heading 2 subsection")
    require(headings[0][0] <= 2, f"document starts at Heading {headings[0][0]}")

    previous_level = None
    for level, text in headings:
        if level == 1:
            require(CHAPTER_TITLE_RE.match(text) is not None, f"Heading 1 is not a Chinese numbered chapter: {text}")
        if previous_level is not None:
            require(level <= previous_level + 1, f"heading level skips from {previous_level} to {level}: {text}")
        previous_level = level

    return {
        "heading_1": sum(level == 1 for level, _ in headings),
        "heading_2": sum(level == 2 for level, _ in headings),
        "heading_3": sum(level == 3 for level, _ in headings),
    }


def audit(path: Path, tokens_path: Path):
    tokens = read_tokens(tokens_path)
    page = tokens["page"]
    fonts = tokens["fonts"]
    colors = tokens["colors"]
    sizes = tokens["sizes_pt"]
    doc = Document(path)
    require(len(doc.sections) == 1, f"sections={len(doc.sections)}")
    section = doc.sections[0]
    require(section.page_width == Twips(page["width_dxa"]), f"page_width={section.page_width}")
    require(section.page_height == Twips(page["height_dxa"]), f"page_height={section.page_height}")
    for name, expected in (("top_margin", page["top_margin_dxa"]), ("bottom_margin", page["bottom_margin_dxa"]), ("left_margin", page["left_margin_dxa"]), ("right_margin", page["right_margin_dxa"])):
        require(getattr(section, name) == Twips(expected), f"{name}={getattr(section, name)}")
    require(section.different_first_page_header_footer, "first-page header/footer must be different")

    style_checks = {
        "Normal": (fonts["body"], sizes["body"]),
        "Heading 1": (fonts["heading"], sizes["h1"]),
        "Heading 2": (fonts["heading"], sizes["h2"]),
        "Heading 3": (fonts["heading"], sizes["h3"]),
    }
    for name, (font, size) in style_checks.items():
        style = doc.styles[name]
        require(style.font.size.pt == size, f"{name}.size={style.font.size.pt}")
        style_xml = style._element.xml
        require(font in style_xml, f"{name} missing font {font}")
    heading_counts = audit_heading_hierarchy(doc)

    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        header_xml = archive.read("word/header1.xml").decode("utf-8") if "word/header1.xml" in archive.namelist() else ""
        footer_xml = archive.read("word/footer1.xml").decode("utf-8") if "word/footer1.xml" in archive.namelist() else ""
        package_xml = "\n".join(archive.read(name).decode("utf-8", "ignore") for name in archive.namelist() if name.endswith(".xml"))
    require(colors["ink"] in document_xml, "missing charcoal ink")
    require(colors["gold"] in package_xml, "missing copper gold")
    require(colors["table_header_fill"] if "table_header_fill" in colors else colors["ink"] in document_xml, "missing table header fill")
    require(fonts["body"] in styles_xml or fonts["body"] in document_xml, "missing body font")
    require(fonts["heading"] in styles_xml or fonts["heading"] in document_xml, "missing heading font")
    require("{{" not in document_xml and "}}" not in document_xml, "unresolved template placeholder")
    require(r"\n" not in document_xml and r"\r" not in document_xml, "escaped line-break marker rendered literally")
    require("PAGE" in footer_xml and "NUMPAGES" in footer_xml, "missing page fields")
    rendered_parts = document_xml + header_xml + footer_xml
    require("2E74B5" not in rendered_parts and "4F81BD" not in rendered_parts, "legacy blue palette remains in rendered parts")

    require(len(doc.tables) >= 3, f"tables={len(doc.tables)}")
    for index, table in enumerate(doc.tables):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find("w:tblW", NS)
        tbl_ind = tbl_pr.find("w:tblInd", NS)
        require(tbl_w is not None, f"table {index} missing tblW")
        require(int(tbl_w.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w")) == page["usable_width_dxa"], f"table {index} width")
        require(tbl_ind is not None and int(tbl_ind.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w")) == page["table_indent_dxa"], f"table {index} indent")
        if not table.rows:
            continue
        first_row = table.rows[0]
        fills = []
        for cell in first_row.cells:
            shd = cell._tc.get_or_add_tcPr().find("w:shd", NS)
            fills.append(None if shd is None else shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill"))
        expected_header_color = colors["gold_on_dark"] if index == 0 else colors["white"] if colors["ink"] in fills else None
        if expected_header_color:
            for cell in first_row.cells:
                visible_runs = [run for paragraph in cell.paragraphs for run in paragraph.runs if run.text]
                require(visible_runs, f"table {index} header cell has no visible text run")
                for run in visible_runs:
                    actual = str(run.font.color.rgb or "").upper()
                    require(actual == expected_header_color, f"table {index} header text color={actual}, expected {expected_header_color}")

    text_parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.extend(p.text for p in cell.paragraphs)
    text = "\n".join(text_parts)
    properties = doc.core_properties
    title = (properties.title or "").strip()
    author = (properties.author or "").strip()
    subject = (properties.subject or "").strip()
    require(title, "missing document title metadata")
    require(author == "财经内容台", f"document author={author!r}, expected '财经内容台'")
    require(subject, "missing document subject metadata")
    require(title in text, "document title metadata is not reflected in document content")
    require("财经内容台" in text, "missing cover organization")
    require("投资要点" in text, "missing investment callout")
    require("关键数字" in text, "missing key metrics label")
    return {
        "status": "PASS",
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "headings": heading_counts,
        "metadata": {"title": title, "author": author, "subject": subject},
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("--tokens", type=Path, default=Path(__file__).resolve().parent.parent / "references" / "docx-style-tokens.json")
    args = parser.parse_args()
    print(audit(args.input, args.tokens))


if __name__ == "__main__":
    main()
